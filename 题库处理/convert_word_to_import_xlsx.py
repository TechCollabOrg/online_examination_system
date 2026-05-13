import argparse
import re
import shutil
import subprocess
import tempfile
import mimetypes
import os
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Set, Tuple

import openpyxl
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from minio import Minio
from minio.error import S3Error


@dataclass
class ParsedQuestion:
    qtype: int  # 1单选 2多选 3判断 4简答
    title: str
    options: List[str]  # A-F
    correct_letters: List[str]  # e.g. ["A"] or ["A","C"]
    answer_text: Optional[str] = None  # for 简答题/无法识别时
    analysis: Optional[str] = None
    stem_image_links: List[str] = None
    option_image_links: List[List[str]] = None  # 6 lists, A-F
    # 模板无独立列：写入第 9 列「解析」文本中的附图链接
    answer_image_links: List[str] = None
    analysis_image_links: List[str] = None


_QUESTION_START_RE = re.compile(r"^\s*(\d{1,4})\s*[\.、．]\s*(.+?)\s*$")
# 一些资料会写成“（1）xxx / (1)xxx / 第1题xxx / 问题1：xxx”
_QUESTION_START_ALT_RE = re.compile(
    r"^\s*(?:[（(]\s*(\d{1,4})\s*[）)]\s*|第\s*(\d{1,4})\s*题\s*|问题\s*(\d{1,4})\s*[:：]?\s*)(.+?)\s*$"
)
# 下午题/案例分析常见写法：
_BIG_Q_RE = re.compile(r"^\s*题目\s*(\d{1,4})\s*[、:：]\s*(.+?)\s*$")
# 【问题1】或 [问题1]（半角）、全角［］；后面可跟 (9分) 等，题干常在下一行
_SUB_Q_RE = re.compile(
    r"^\s*(?:【|［|\[)\s*问题\s*(\d{1,4})\s*(?:】|］|\])\s*(.*)$"
)
# 如「试题一」「试题2」单独成行，作为大题分界（与「题目1、」并存）
_SHITI_RE = re.compile(r"^\s*试题\s*([一二三四五六七八九十百千万两〇零0-9]{1,6})\s*$")

_OPTION_RE = re.compile(r"^\s*([A-F])\s*[\.\、．\)\]]\s*(.+?)\s*$", re.IGNORECASE)

# 常见“答案/解析”写法（尽量宽松，覆盖真题/模拟题排版）
_ANSWER_LINE_RE = re.compile(
    r"^\s*(?:"
    r"答案|参考答案|正确答案|标准答案|试题答案|本题答案|"
    r"【\s*答案\s*】|【答案】|"
    r"答\s*[:：]|选\s*[:：]|应选|答案为|正确选项|"
    r"第\s*\d{1,4}\s*题\s*答案"
    r")\s*[:：]?\s*(.+?)\s*$"
)
_ANALYSIS_LINE_RE = re.compile(
    r"^\s*(?:解析|【\s*解析\s*】|【解析】|答案解析|试题解析|详解)\s*[:：]?\s*(.*)$"
)
# 整行仅为选项字母：如 "D" 或 "A、C" 或 "AC"（在已有选项的题块末尾常见）
_STANDALONE_ANSWER_LETTERS_RE = re.compile(
    r"^\s*([A-F](?:\s*[,，、]\s*[A-F])*|[A-F]{2,6})\s*$", re.IGNORECASE
)


def _normalize_text(s: str) -> str:
    s = s.replace("\u00a0", " ")
    s = s.replace("\r\n", "\n").replace("\r", "\n")
    # 合并重复空格，但保留换行（便于规则解析）
    s = re.sub(r"[ \t]+", " ", s)
    return s.strip()


def _strip_img_markers(text: str) -> Tuple[str, List[str]]:
    """从文本中取出 [[IMG:路径]] 标记，返回（干净题干, 图片路径列表）。"""
    paths = re.findall(r"\[\[IMG:(.+?)\]\]", text)
    clean = re.sub(r"\s*\[\[IMG:.+?\]\]\s*", "\n", text)
    clean = re.sub(r"\n{3,}", "\n\n", clean).strip()
    return clean, paths


def extract_text_from_docx(docx_path: Path) -> str:
    doc = Document(str(docx_path))
    parts: List[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    return _normalize_text("\n".join(parts))


def _safe_slug(name: str) -> str:
    # 用于生成文件夹/文件名：尽量保留中文，但去掉不安全字符
    name = re.sub(r"[<>:\"/\\\\|?*]+", "_", name)
    name = name.strip().strip(".")
    return name or "paper"


def _iter_cell_paragraphs(cell) -> List:
    """单元格内段落 + 嵌套表格（递归）。"""
    out: List = []
    for p in cell.paragraphs:
        out.append(p)
    for tbl in cell.tables:
        for row in tbl.rows:
            for c in row.cells:
                out.extend(_iter_cell_paragraphs(c))
    return out


def iter_all_paragraphs_in_order(doc: Document) -> List:
    """按 Word 文档顺序：正文块（段落 / 表）中的全部段落（含表格内）。"""
    paras: List = []
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            paras.append(Paragraph(child, doc))
        elif child.tag == qn("w:tbl"):
            table = Table(child, doc)
            for row in table.rows:
                for cell in row.cells:
                    paras.extend(_iter_cell_paragraphs(cell))
    return paras


def extract_text_and_images_from_docx(docx_path: Path, assets_dir: Path) -> str:
    """
    把 docx 转成“带图片标记的纯文本”。
    图片会导出到 assets_dir，并在文本中用 [[IMG:绝对路径]] 标记。

    目标：让后续解析能把图片链接绑定到题干/选项。
    """
    doc = Document(str(docx_path))
    assets_dir.mkdir(parents=True, exist_ok=True)

    # 建立 rId -> 图片文件路径 的映射
    rid_to_path: Dict[str, str] = {}
    img_idx = 0
    for rid, rel in doc.part.rels.items():
        if rel.reltype != RT.IMAGE:
            continue
        img_idx += 1
        part = rel.target_part
        ext = Path(part.partname).suffix or ".png"
        out_path = assets_dir / f"img_{img_idx:04d}{ext}"
        with open(out_path, "wb") as f:
            f.write(part.blob)
        rid_to_path[rid] = str(out_path)

    def paragraph_to_text_with_imgs(p) -> str:
        # 如果段落包含图片（drawing），把对应 rId 写成 IMG 标记
        embeds = []
        try:
            # 注意：python-docx 的 oxml 元素 .xpath 不支持 namespaces 参数。
            # 这里用 local-name() 避免命名空间问题。
            blips = p._p.xpath(".//*[local-name()='blip']")  # type: ignore[attr-defined]
            for b in blips:
                rid = None
                # 优先找 r:embed（命名空间展开后的 key）
                for k, v in b.attrib.items():
                    if k.endswith("}embed") or k == "embed":
                        rid = v
                        break
                if rid and rid in rid_to_path:
                    embeds.append(rid_to_path[rid])
        except Exception:
            embeds = []

        t = (p.text or "").strip()
        if not t and not embeds:
            return ""
        chunks = []
        if t:
            chunks.append(t)
        for img in embeds:
            chunks.append(f"[[IMG:{img}]]")
        return " ".join(chunks).strip()

    parts: List[str] = []
    for p in iter_all_paragraphs_in_order(doc):
        line = paragraph_to_text_with_imgs(p)
        if line:
            parts.append(line)

    return _normalize_text("\n".join(parts))


def _normalize_minio_endpoint(endpoint: str) -> Tuple[str, bool]:
    endpoint = endpoint.strip()
    if endpoint.startswith("http://"):
        return endpoint[len("http://") :], False
    if endpoint.startswith("https://"):
        return endpoint[len("https://") :], True
    # 默认按 http
    return endpoint, False


def _collect_used_image_paths(questions: List[ParsedQuestion]) -> Set[Path]:
    used: Set[Path] = set()
    for q in questions:
        for s in (q.stem_image_links or []):
            try:
                used.add(Path(s))
            except Exception:
                continue
        for s in (q.answer_image_links or []):
            try:
                used.add(Path(s))
            except Exception:
                continue
        for s in (q.analysis_image_links or []):
            try:
                used.add(Path(s))
            except Exception:
                continue
        for opt_list in (q.option_image_links or [[] for _ in range(6)]):
            for s in opt_list:
                try:
                    used.add(Path(s))
                except Exception:
                    continue
    return used


def upload_used_images_to_minio(
    *,
    questions: List[ParsedQuestion],
    endpoint: str,
    access_key: str,
    secret_key: str,
    bucket: str,
    object_prefix: str,
    public_base_url: Optional[str],
    presign_seconds: int,
) -> None:
    """
    只上传“已绑定到题干/选项”的图片，并把 questions 里的本地路径替换为 MinIO 链接。

    - public_base_url：如果你的 MinIO 对外访问地址与 endpoint 不同（例如经 Nginx 反代），传这个。
    - presign_seconds > 0：生成带签名的临时访问链接（适合私有桶）。
      presign_seconds = 0：直接拼 URL（适合桶/网关公开访问）。
    """
    host, secure = _normalize_minio_endpoint(endpoint)
    client = Minio(host, access_key=access_key, secret_key=secret_key, secure=secure)

    # 确保桶存在
    try:
        if not client.bucket_exists(bucket):
            client.make_bucket(bucket)
    except S3Error as e:
        raise RuntimeError(f"MinIO 桶检查/创建失败：{e}") from e

    used_paths = _collect_used_image_paths(questions)
    if not used_paths:
        return

    # 上传并建立映射：本地路径 -> url
    mapping: Dict[str, str] = {}
    prefix = object_prefix.strip().strip("/")
    for p in sorted(used_paths):
        if not p.exists() or not p.is_file():
            continue
        obj_name = f"{prefix}/{p.parent.name}/{p.name}" if prefix else f"{p.parent.name}/{p.name}"
        obj_name = obj_name.replace("\\", "/")
        ctype, _ = mimetypes.guess_type(str(p))
        try:
            client.fput_object(bucket, obj_name, str(p), content_type=ctype or "application/octet-stream")
        except S3Error as e:
            raise RuntimeError(f"MinIO 上传失败：{p} -> {e}") from e

        if presign_seconds > 0:
            url = client.presigned_get_object(bucket, obj_name, expires=presign_seconds)
        else:
            base = (public_base_url or endpoint).rstrip("/")
            url = f"{base}/{bucket}/{obj_name}"
        mapping[str(p)] = url

    def repl_one(s: str) -> str:
        return mapping.get(s, s)

    for q in questions:
        q.stem_image_links = [repl_one(x) for x in (q.stem_image_links or [])]
        q.option_image_links = [
            [repl_one(x) for x in lst] for lst in (q.option_image_links or [[] for _ in range(6)])
        ]


def _which(cmd: str) -> Optional[str]:
    return shutil.which(cmd)


def convert_doc_to_docx_with_soffice(doc_path: Path, out_dir: Path) -> Optional[Path]:
    """
    用 LibreOffice/soffice 把 .doc 转 .docx（优先方案，避免依赖 MS Word）。
    """
    soffice = _which("soffice") or _which("soffice.exe")
    if not soffice:
        return None

    out_dir.mkdir(parents=True, exist_ok=True)
    # soffice 会在 out_dir 里生成同名 .docx
    cmd = [
        soffice,
        "--headless",
        "--nologo",
        "--nofirststartwizard",
        "--convert-to",
        "docx",
        "--outdir",
        str(out_dir),
        str(doc_path),
    ]
    try:
        subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    except Exception:
        return None

    expected = out_dir / (doc_path.stem + ".docx")
    return expected if expected.exists() else None


def extract_text_from_word(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".docx":
        return extract_text_from_docx(path)
    if ext == ".doc":
        with tempfile.TemporaryDirectory(prefix="word2xlsx_") as td:
            out_dir = Path(td)
            docx = convert_doc_to_docx_with_soffice(path, out_dir)
            if docx:
                return extract_text_from_docx(docx)
        raise RuntimeError(
            "无法读取 .doc：未检测到 LibreOffice（soffice），或转换失败。"
            "请安装 LibreOffice 并确保 soffice 在 PATH 中。"
        )
    raise ValueError(f"不支持的文件类型：{path}")


def _split_into_question_blocks(lines: List[str]) -> List[List[str]]:
    blocks: List[List[str]] = []
    current: List[str] = []

    for line in lines:
        if (
            _QUESTION_START_RE.match(line)
            or _QUESTION_START_ALT_RE.match(line)
            or _BIG_Q_RE.match(line)
            or _SUB_Q_RE.match(line)
            or _SHITI_RE.match(line)
        ):
            if current:
                blocks.append(current)
            current = [line]
        else:
            if current:
                current.append(line)
            else:
                # 文档开头的标题/说明，先忽略
                continue

    if current:
        blocks.append(current)
    return blocks


def _parse_correct_letters(s: str) -> List[str]:
    s = s.strip()
    # 形如：A / AC / A,C / A C / A、C
    letters = re.findall(r"[A-F]", s.upper())
    if letters:
        # 去重并保持顺序
        seen = set()
        out: List[str] = []
        for x in letters:
            if x not in seen:
                seen.add(x)
                out.append(x)
        return out

    # 判断题：对/错/正确/错误
    if any(k in s for k in ["对", "正确", "是", "√"]):
        return ["A"]  # A=正确
    if any(k in s for k in ["错", "错误", "否", "×"]):
        return ["B"]  # B=错误

    return []


def parse_questions_from_text(text: str) -> List[ParsedQuestion]:
    lines = [x.strip() for x in text.split("\n") if x.strip()]
    # 先处理“题目X + 【问题Y】”结构：把每个【问题】拆成独立题（简答题）
    # 如果文档没有【问题】结构，再走普通的“1. / A. / 答案：”解析。
    has_subq = any(_SUB_Q_RE.match(x) for x in lines)
    if has_subq:
        parsed: List[ParsedQuestion] = []
        current_big: Optional[Tuple[str, str]] = None  # ("题目1", "题目说明行")
        current_q_title: Optional[str] = None
        current_body: List[str] = []

        def flush_subq() -> None:
            nonlocal current_q_title, current_body
            if not current_q_title:
                return
            answer_text: Optional[str] = None
            analysis_lines_sub: List[str] = []
            body_remain: List[str] = []
            answer_imgs_sq: List[str] = []
            analysis_imgs_sq: List[str] = []
            stem_pre_imgs: List[str] = []
            in_an = False
            after_ans = False

            for line in current_body:
                line_work = line
                if "[[IMG:" in line:
                    imgs = re.findall(r"\[\[IMG:(.+?)\]\]", line)
                    line_work = re.sub(r"\s*\[\[IMG:.+?\]\]\s*", " ", line).strip()
                    if in_an:
                        analysis_imgs_sq.extend(imgs)
                    elif after_ans:
                        answer_imgs_sq.extend(imgs)
                    else:
                        stem_pre_imgs.extend(imgs)
                    if not line_work:
                        continue

                am = _ANSWER_LINE_RE.match(line_work)
                if am:
                    after_ans = True
                    answer_text = am.group(1).strip()
                    continue
                anm = _ANALYSIS_LINE_RE.match(line_work)
                if anm:
                    in_an = True
                    after_ans = False
                    rest = (anm.group(1) or "").strip()
                    if rest:
                        analysis_lines_sub.append(rest)
                    continue
                if in_an:
                    analysis_lines_sub.append(line_work)
                    continue
                if after_ans:
                    answer_text = (answer_text or "") + ("\n" + line_work if answer_text else line_work)
                    continue
                body_remain.append(line_work)

            body = "\n".join(body_remain).strip()
            title_full = current_q_title
            if body:
                title_full = title_full + "\n" + body
            analysis_sub = "\n".join([x for x in analysis_lines_sub if x]).strip() or None
            title_full, stem_imgs_sq = _strip_img_markers(title_full.strip())
            stem_imgs_sq = stem_pre_imgs + stem_imgs_sq
            parsed.append(
                ParsedQuestion(
                    qtype=4,
                    title=title_full,
                    options=[""] * 6,
                    correct_letters=[],
                    answer_text=answer_text,
                    analysis=analysis_sub,
                    stem_image_links=stem_imgs_sq,
                    option_image_links=[[] for _ in range(6)],
                    answer_image_links=answer_imgs_sq,
                    analysis_image_links=analysis_imgs_sq,
                )
            )
            current_q_title = None
            current_body = []

        for line in lines:
            es = _SHITI_RE.match(line)
            if es:
                flush_subq()
                current_big = (f"试题{es.group(1)}", "")
                continue

            bm = _BIG_Q_RE.match(line)
            if bm:
                flush_subq()
                current_big = (f"题目{bm.group(1)}", bm.group(2).strip())
                continue

            sm = _SUB_Q_RE.match(line)
            if sm:
                flush_subq()
                prefix = ""
                if current_big:
                    big_no, big_title = current_big
                    if big_title:
                        prefix = f"{big_no}：{big_title}\n"
                    else:
                        prefix = f"{big_no}\n"
                rest = (sm.group(2) or "").strip()
                current_q_title = f"{prefix}问题{sm.group(1)}".strip()
                if rest:
                    current_q_title = f"{current_q_title}：{rest}"
                continue

            if current_q_title:
                current_body.append(line)

        flush_subq()
        return parsed

    blocks = _split_into_question_blocks(lines)
    parsed: List[ParsedQuestion] = []

    for b in blocks:
        title: Optional[str] = None
        m = _QUESTION_START_RE.match(b[0])
        if m:
            title = m.group(2).strip()
        else:
            am = _QUESTION_START_ALT_RE.match(b[0])
            if am:
                title = am.group(4).strip()
            else:
                bm = _BIG_Q_RE.match(b[0])
                if bm:
                    title = f"题目{bm.group(1)}：{bm.group(2).strip()}"
                else:
                    sh = _SHITI_RE.match(b[0])
                    if sh:
                        title = f"试题{sh.group(1)}"
                    else:
                        sm = _SUB_Q_RE.match(b[0])
                        if sm:
                            rest = (sm.group(2) or "").strip()
                            title = f"问题{sm.group(1)}" + (f"：{rest}" if rest else "")
        if not title:
            continue
        options: List[str] = []
        correct_letters: List[str] = []
        answer_text: Optional[str] = None
        analysis_lines: List[str] = []
        stem_imgs: List[str] = []
        option_imgs: List[List[str]] = [[] for _ in range(6)]
        answer_imgs: List[str] = []
        analysis_imgs: List[str] = []

        in_analysis = False
        after_answer = False
        current_option_idx: Optional[int] = None

        for line in b[1:]:
            # 图片标记：按阶段绑定到题干 / 选项 / 答案区 / 解析区
            if "[[IMG:" in line:
                imgs = re.findall(r"\[\[IMG:(.+?)\]\]", line)
                # 图片标记可能和文本在同一行：先把标记去掉再继续解析文本
                line_wo = re.sub(r"\s*\[\[IMG:.+?\]\]\s*", " ", line).strip()
                if in_analysis:
                    analysis_imgs.extend(imgs)
                elif after_answer:
                    answer_imgs.extend(imgs)
                elif current_option_idx is not None:
                    option_imgs[current_option_idx].extend(imgs)
                else:
                    stem_imgs.extend(imgs)
                if not line_wo:
                    continue
                line = line_wo

            optm = _OPTION_RE.match(line)
            if optm and not in_analysis:
                raw_line = line
                embedded = re.findall(r"\[\[IMG:(.+?)\]\]", raw_line)
                line_clean = re.sub(r"\s*\[\[IMG:.+?\]\]\s*", " ", raw_line).strip()
                optm = _OPTION_RE.match(line_clean)
                if not optm:
                    continue
                letter = optm.group(1).upper()
                content = optm.group(2).strip()
                idx = ord(letter) - ord("A")
                while len(options) <= idx:
                    options.append("")
                options[idx] = content
                option_imgs[idx].extend(embedded)
                current_option_idx = None
                continue

            am = _ANSWER_LINE_RE.match(line)
            if am:
                val = am.group(1).strip()
                correct_letters = _parse_correct_letters(val)
                after_answer = True
                current_option_idx = None
                if not correct_letters:
                    answer_text = val
                continue

            sam = _STANDALONE_ANSWER_LETTERS_RE.match(line)
            if (
                sam
                and not in_analysis
                and not correct_letters
                and any(options)
            ):
                letters = _parse_correct_letters(sam.group(1))
                if letters:
                    correct_letters = letters
                    after_answer = True
                    current_option_idx = None
                    continue

            anm = _ANALYSIS_LINE_RE.match(line)
            if anm:
                in_analysis = True
                after_answer = False
                rest = (anm.group(1) or "").strip()
                if rest:
                    analysis_lines.append(rest)
                continue

            if in_analysis:
                analysis_lines.append(line.strip())
            else:
                # 有些题干会换行续写
                if not _OPTION_RE.match(line) and not _ANSWER_LINE_RE.match(line):
                    # 题干追加（避免把“（1）…”这类当成新题）
                    title += "\n" + line.strip()
                    current_option_idx = None

        analysis = "\n".join([x for x in analysis_lines if x]).strip() or None

        # 推断题型
        if options:
            qtype = 2 if len(correct_letters) >= 2 else 1
        else:
            # 没有选项时，优先当简答；但如果答案能识别为“对/错”则当判断
            if correct_letters in (["A"], ["B"]) and (answer_text is None):
                qtype = 3
            else:
                qtype = 4

        title_clean, title_imgs = _strip_img_markers(title.strip())
        stem_imgs = title_imgs + stem_imgs

        parsed.append(
            ParsedQuestion(
                qtype=qtype,
                title=title_clean,
                options=(options + [""] * 6)[:6],
                correct_letters=correct_letters,
                answer_text=answer_text,
                analysis=analysis,
                stem_image_links=stem_imgs,
                option_image_links=option_imgs,
                answer_image_links=answer_imgs,
                analysis_image_links=analysis_imgs,
            )
        )

    return parsed


def write_questions_to_template(
    template_path: Path, out_path: Path, questions: List[ParsedQuestion]
) -> None:
    wb = openpyxl.load_workbook(str(template_path))
    ws = wb.active

    # 模板：第1行“填写须知”，第2行“列名”，第3行起是示例，先清空示例行（保留 1-2 行）
    if ws.max_row >= 3:
        ws.delete_rows(3, ws.max_row - 2)

    def set_row(row_idx: int, q: ParsedQuestion) -> None:
        # 约定列（与模板示例行一致）：
        # 1:题目类型(1-4) 2:题干
        # 3-8:选项A-F内容
        # 9:解析
        # 10-15:选项A-F是否正确(1/0)
        # 16:题干图片 17-22:选项A-F图片（这里先留空）
        # 23-25:材料组编号 / 共用材料题干 / 共用材料题干图片（导入多小问时手填；本脚本默认不写，留空即可）
        ws.cell(row=row_idx, column=1, value=q.qtype)
        ws.cell(row=row_idx, column=2, value=q.title)

        # 判断题特殊处理：写成 A=正确，B=错误，便于模板导入
        options = list(q.options)
        correct_letters = list(q.correct_letters)
        if q.qtype == 3:
            options = ["正确", "错误", "", "", "", ""]
            if not correct_letters and q.answer_text:
                correct_letters = _parse_correct_letters(q.answer_text)

        for i in range(6):
            ws.cell(row=row_idx, column=3 + i, value=(options[i] or None))

        analysis_parts: List[str] = []
        if q.answer_text:
            analysis_parts.append(f"答案：{q.answer_text}")
        elif q.correct_letters and q.qtype in (1, 2, 3):
            analysis_parts.append("答案：" + "、".join(q.correct_letters))
        if q.answer_image_links:
            analysis_parts.append("答案附图：\n" + "\n".join(q.answer_image_links))
        if q.analysis:
            analysis_parts.append(q.analysis)
        if q.analysis_image_links:
            analysis_parts.append("解析附图：\n" + "\n".join(q.analysis_image_links))
        ws.cell(row=row_idx, column=9, value="\n".join(analysis_parts).strip() or None)

        correct_set = set([x.upper() for x in correct_letters])
        for i, letter in enumerate(["A", "B", "C", "D", "E", "F"]):
            val = 1 if letter in correct_set else 0
            # 如果该选项为空，就不写“是否正确”，避免导入误判
            if options[i]:
                ws.cell(row=row_idx, column=10 + i, value=val)
            else:
                ws.cell(row=row_idx, column=10 + i, value=None)

        # 图片列留空
        stem_links = (q.stem_image_links or [])
        opt_links = (q.option_image_links or [[] for _ in range(6)])
        ws.cell(row=row_idx, column=16, value="\n".join(stem_links) or None)
        for i in range(6):
            ws.cell(row=row_idx, column=17 + i, value="\n".join(opt_links[i]) or None)

    for idx, q in enumerate(questions, start=3):
        set_row(idx, q)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(out_path))


def validate_questions(questions: List[ParsedQuestion]) -> List[str]:
    issues: List[str] = []
    for i, q in enumerate(questions, start=1):
        if not q.title.strip():
            issues.append(f"第{i}题：题干为空")
        if q.qtype in (1, 2):
            # 至少要有 A/B 选项
            if not (q.options and len(q.options) >= 2 and q.options[0] and q.options[1]):
                issues.append(f"第{i}题：选择题缺少 A/B 选项")
            if not q.correct_letters:
                issues.append(f"第{i}题：选择题未识别到正确答案（答案行可能不规范）")
        if q.qtype == 3:
            # 判断题必须有 A 或 B
            if q.correct_letters and not set(q.correct_letters).issubset({"A", "B"}):
                issues.append(f"第{i}题：判断题答案不应出现 A/B 之外的字母：{q.correct_letters}")
        # 简答题允许没有 answer/analysis
    return issues


def collect_input_files(inputs: List[str]) -> List[Path]:
    out: List[Path] = []
    for s in inputs:
        p = Path(s)
        if p.is_dir():
            for ext in (".doc", ".docx"):
                out.extend(sorted(p.rglob(f"*{ext}")))
        else:
            out.append(p)
    # 去重、保序
    seen = set()
    unique: List[Path] = []
    for p in out:
        rp = str(p.resolve())
        if rp not in seen:
            seen.add(rp)
            unique.append(p)
    return unique


def main() -> int:
    ap = argparse.ArgumentParser(
        description="把 .doc/.docx 题库批量转换为“导入试题模板”同结构 .xlsx"
    )
    ap.add_argument(
        "--template",
        default=str(
            Path(__file__).resolve().parents[1]
            / "online-exam-system-frontend/public/template/ImportQuestionTemplate.xlsx"
        ),
        help="导入试题模板 .xlsx 路径（默认用项目自带模板）",
    )
    ap.add_argument(
        "--out-dir",
        default=str(Path.cwd() / "out"),
        help="输出目录（默认 ./out）",
    )
    ap.add_argument(
        "--per-paper",
        action="store_true",
        help="按每个输入文件分别生成一个 .xlsx（推荐，便于按卷子整理）",
    )
    ap.add_argument(
        "--strict",
        action="store_true",
        help="严格模式：若发现明显无法导入的问题（如选择题无答案/无A/B选项）则返回非0，并生成报告",
    )
    ap.add_argument("--minio-endpoint", default=os.environ.get("MINIO_ENDPOINT", ""), help="MinIO 地址，例如 http://127.0.0.1:9000")
    # 兼容两种常见写法：MINIO_ACCESS_KEY / MINIO_SECRET_KEY 以及 MINIO_ACCESSKEY / MINIO_SECRETKEY
    ap.add_argument(
        "--minio-access-key",
        default=os.environ.get("MINIO_ACCESS_KEY", "") or os.environ.get("MINIO_ACCESSKEY", ""),
        help="MinIO AccessKey（也可用环境变量 MINIO_ACCESS_KEY / MINIO_ACCESSKEY）",
    )
    ap.add_argument(
        "--minio-secret-key",
        default=os.environ.get("MINIO_SECRET_KEY", "") or os.environ.get("MINIO_SECRETKEY", ""),
        help="MinIO SecretKey（也可用环境变量 MINIO_SECRET_KEY / MINIO_SECRETKEY）",
    )
    ap.add_argument("--minio-bucket", default=os.environ.get("MINIO_BUCKET", "online-exam"), help="MinIO 桶名（默认 online-exam）")
    ap.add_argument("--minio-prefix", default="exam-files/questions", help="对象名前缀（默认 exam-files/questions）")
    ap.add_argument(
        "--minio-public-base-url",
        default=os.environ.get("MINIO_PUBLIC_BASE_URL", ""),
        help="对外访问的 MinIO Base URL（可选；若不填则使用 --minio-endpoint）",
    )
    ap.add_argument(
        "--minio-presign-seconds",
        type=int,
        default=0,
        help="生成临时访问链接的有效期秒数（0=不签名直接拼 URL；私有桶建议设置如 604800=7天）",
    )
    ap.add_argument(
        "inputs",
        nargs="+",
        help="输入文件或文件夹（支持 .doc/.docx；文件夹会递归扫描）",
    )

    args = ap.parse_args()
    template_path = Path(args.template)
    out_dir = Path(args.out_dir)

    if not template_path.exists():
        raise FileNotFoundError(f"模板不存在：{template_path}")

    files = collect_input_files(args.inputs)
    if not files:
        print("没有找到任何 .doc/.docx")
        return 2

    report_lines: List[str] = []
    strict_failed = False
    all_questions: List[ParsedQuestion] = []
    for f in files:
        try:
            # 对 docx：导出图片并插入 IMG 标记，后续写入模板图片链接列
            if f.suffix.lower() == ".docx":
                paper_slug = _safe_slug(f.stem)
                assets_dir = out_dir / "assets" / paper_slug
                text = extract_text_and_images_from_docx(f, assets_dir)
            else:
                text = extract_text_from_word(f)
            qs = parse_questions_from_text(text)
            if not qs:
                print(f"[WARN] 未识别到题目：{f}")
                report_lines.append(f"[WARN] 未识别到题目：{f}")
            else:
                print(f"[OK] {f} -> 识别 {len(qs)} 题")
            issues = validate_questions(qs)
            if issues:
                report_lines.append(f"[CHECK] {f} 发现 {len(issues)} 个可能影响导入的问题：")
                report_lines.extend([f"  - {x}" for x in issues])
                if args.strict:
                    strict_failed = True
            if args.per_paper:
                ts = datetime.now().strftime("%Y%m%d-%H%M%S")
                out_path = out_dir / f"{_safe_slug(f.stem)}-{ts}.xlsx"

                # 如果配置了 MinIO，就把“已绑定的图片”上传并替换为 MinIO 链接（失败则保留本地路径，仍写出 xlsx）
                if args.minio_endpoint:
                    if not (args.minio_access_key and args.minio_secret_key):
                        raise RuntimeError("已设置 --minio-endpoint，但缺少 MinIO 密钥（--minio-access-key/--minio-secret-key）")
                    try:
                        upload_used_images_to_minio(
                            questions=qs,
                            endpoint=args.minio_endpoint,
                            access_key=args.minio_access_key,
                            secret_key=args.minio_secret_key,
                            bucket=args.minio_bucket,
                            object_prefix=args.minio_prefix,
                            public_base_url=args.minio_public_base_url or None,
                            presign_seconds=args.minio_presign_seconds,
                        )
                    except Exception as e:
                        msg = f"[WARN] MinIO 上传失败，图片列将使用本地路径：{f} -> {e}"
                        print(msg)
                        report_lines.append(msg)

                write_questions_to_template(template_path, out_path, qs)
                print(f"已生成：{out_path}（本卷 {len(qs)} 题）")
            else:
                all_questions.extend(qs)
        except Exception as e:
            print(f"[ERR] 处理失败：{f} -> {e}")
            report_lines.append(f"[ERR] 处理失败：{f} -> {e}")

    if not args.per_paper:
        if not all_questions:
            print("没有任何可写入的题目，已退出。")
            return 3

        ts = datetime.now().strftime("%Y%m%d-%H%M%S")
        out_path = out_dir / f"导入题目-{ts}.xlsx"
        write_questions_to_template(template_path, out_path, all_questions)
        print(f"已生成：{out_path}（共 {len(all_questions)} 题）")

    if report_lines:
        ts = datetime.now().strftime("%Y%m%d-%H%M%S")
        report_path = out_dir / f"转换报告-{ts}.txt"
        report_path.parent.mkdir(parents=True, exist_ok=True)
        report_path.write_text("\n".join(report_lines), encoding="utf-8")
        print(f"已生成报告：{report_path}")
    return 4 if strict_failed else 0


if __name__ == "__main__":
    raise SystemExit(main())

